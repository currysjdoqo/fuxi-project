"""
习题导出工具模块
支持导出为 Word (.docx) 和 PDF 格式
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os


def normalize_answer(answer):
    """标准化答案格式"""
    if not answer:
        return ""
    # 移除多余空格
    answer = str(answer).strip()
    return answer


def format_question_for_export(question, include_answer=False, include_analysis=False):
    """格式化题目内容用于导出"""
    result = {
        'type': question.get('type', 'unknown'),
        'content': question.get('content', ''),
        'options': [],
        'answer': '',
        'analysis': ''
    }
    
    # 处理选项
    options = question.get('options', [])
    if isinstance(options, dict):
        for key in sorted(options.keys()):
            result['options'].append({
                'key': key,
                'value': options[key]
            })
    elif isinstance(options, list):
        for opt in options:
            if isinstance(opt, dict):
                result['options'].append({
                    'key': opt.get('key', ''),
                    'value': opt.get('value', '')
                })
    
    # 处理答案
    if include_answer:
        result['answer'] = normalize_answer(question.get('answer', ''))
    
    # 处理解析
    if include_analysis:
        result['analysis'] = question.get('analysis', '') or ''
    
    return result


def export_to_word(questions, subject_name, include_answer=False, include_analysis=False):
    """导出为 Word 文档"""
    doc = Document()
    
    # 设置标题
    title = doc.add_heading(f'{subject_name} 练习题', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加导出信息
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info_run = info_para.add_run(f'共 {len(questions)} 题')
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()  # 空行
    
    # 题号计数器
    q_num = 1
    
    for question in questions:
        formatted = format_question_for_export(question, include_answer, include_analysis)
        
        # 题干
        type_labels = {
            'single': '【单选题】',
            'multiple': '【多选题】',
            'judge': '【判断题】',
            'fill': '【填空题】',
            'short_answer': '【简答题】',
            'programming': '【编程题】'
        }
        type_label = type_labels.get(formatted['type'], '【未知题型】')
        
        # 题目内容
        para = doc.add_paragraph()
        para.add_run(f'{q_num}. {type_label} ').bold = True
        para.add_run(formatted['content'])
        
        # 选项（单选题、多选题、判断题）
        if formatted['options']:
            for opt in formatted['options']:
                opt_para = doc.add_paragraph(style='List Bullet')
                opt_para.add_run(f'{opt["key"]}. {opt["value"]}')
        
        # 填空题显示下划线
        elif formatted['type'] == 'fill':
            fill_para = doc.add_paragraph()
            if include_answer and formatted['answer']:
                fill_para.add_run('答案：____________________（' + formatted['answer'] + '）')
            else:
                fill_para.add_run('答案：____________________')
        
        # 简答题显示答题区域
        elif formatted['type'] == 'short_answer':
            answer_para = doc.add_paragraph()
            if include_answer and formatted['answer']:
                answer_para.add_run('参考答案：' + formatted['answer'])
            else:
                answer_para.add_run('答：________________________')
                answer_para = doc.add_paragraph()
                answer_para.add_run('____________________________')
                answer_para = doc.add_paragraph()
                answer_para.add_run('____________________________')
        
        # 编程题显示代码区域
        elif formatted['type'] == 'programming':
            code_para = doc.add_paragraph()
            code_para.add_run('（请在下方作答）')
            for _ in range(8):
                doc.add_paragraph(' ')
        
        # 答案
        if include_answer and formatted['answer']:
            if formatted['type'] not in ['fill', 'short_answer', 'programming']:
                answer_para = doc.add_paragraph()
                answer_run = answer_para.add_run(f'正确答案：{formatted["answer"]}')
                answer_run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
                answer_run.bold = True
        
        # 解析
        if include_analysis and formatted['analysis']:
            analysis_para = doc.add_paragraph()
            analysis_run = analysis_para.add_run(f'【解析】{formatted["analysis"]}')
            analysis_run.font.color.rgb = RGBColor(128, 128, 128)  # 灰色
            analysis_run.font.size = Pt(9)
        
        # 空行分隔
        doc.add_paragraph()
        
        q_num += 1
    
    # 保存文档
    filename = f'{subject_name}_练习题.docx'
    filepath = os.path.join('exports', filename)
    
    # 确保导出目录存在
    os.makedirs('exports', exist_ok=True)
    
    doc.save(filepath)
    return filepath


def export_to_pdf(questions, subject_name, include_answer=False, include_analysis=False):
    """导出为 PDF 文档"""
    # 创建 PDF
    filename = f'{subject_name}_练习题.pdf'
    filepath = os.path.join('exports', filename)
    
    # 确保导出目录存在
    os.makedirs('exports', exist_ok=True)
    
    doc = SimpleDocTemplate(
        filepath,
        pagesize=A4,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=18
    )
    
    # 获取样式
    styles = getSampleStyleSheet()
    
    # 自定义样式
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        alignment=1,  # 居中
        spaceAfter=20
    )
    
    info_style = ParagraphStyle(
        'InfoStyle',
        parent=styles['Normal'],
        fontSize=10,
        alignment=2,  # 右对齐
        textColor='gray'
    )
    
    question_style = ParagraphStyle(
        'QuestionStyle',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=10
    )
    
    option_style = ParagraphStyle(
        'OptionStyle',
        parent=styles['Normal'],
        fontSize=10,
        leftIndent=20,
        spaceAfter=5
    )
    
    answer_style = ParagraphStyle(
        'AnswerStyle',
        parent=styles['Normal'],
        fontSize=10,
        textColor='green',
        spaceAfter=10
    )
    
    analysis_style = ParagraphStyle(
        'AnalysisStyle',
        parent=styles['Normal'],
        fontSize=9,
        textColor='gray',
        spaceAfter=15
    )
    
    # 构建内容
    story = []
    
    # 标题
    story.append(Paragraph(f'{subject_name} 练习题', title_style))
    story.append(Paragraph(f'共 {len(questions)} 题', info_style))
    story.append(Spacer(1, 20))
    
    # 题号计数器
    q_num = 1
    
    for question in questions:
        formatted = format_question_for_export(question, include_answer, include_analysis)
        
        # 题干
        type_labels = {
            'single': '【单选题】',
            'multiple': '【多选题】',
            'judge': '【判断题】',
            'fill': '【填空题】',
            'short_answer': '【简答题】',
            'programming': '【编程题】'
        }
        type_label = type_labels.get(formatted['type'], '【未知题型】')
        
        # 题目内容
        content = f'<b>{q_num}. {type_label}</b> {formatted["content"]}'
        story.append(Paragraph(content, question_style))
        
        # 选项
        if formatted['options']:
            for opt in formatted['options']:
                option_text = f'{opt["key"]}. {opt["value"]}'
                story.append(Paragraph(option_text, option_style))
        
        # 填空题
        elif formatted['type'] == 'fill':
            if include_answer and formatted['answer']:
                story.append(Paragraph(
                    f'答案：____________________（{formatted["answer"]}）',
                    question_style
                ))
            else:
                story.append(Paragraph('答案：____________________', question_style))
        
        # 简答题
        elif formatted['type'] == 'short_answer':
            if include_answer and formatted['answer']:
                story.append(Paragraph(f'参考答案：{formatted["answer"]}', question_style))
            else:
                story.append(Paragraph('答：________________________', question_style))
                story.append(Paragraph('____________________________', question_style))
        
        # 编程题
        elif formatted['type'] == 'programming':
            story.append(Paragraph('（请在下方作答）', question_style))
            for _ in range(6):
                story.append(Paragraph('_' * 60, question_style))
        
        # 答案
        if include_answer and formatted['answer']:
            if formatted['type'] not in ['fill', 'short_answer', 'programming']:
                story.append(Paragraph(
                    f'<b><font color="green">正确答案：{formatted["answer"]}</font></b>',
                    answer_style
                ))
        
        # 解析
        if include_analysis and formatted['analysis']:
            story.append(Paragraph(
                f'<i>【解析】{formatted["analysis"]}</i>',
                analysis_style
            ))
        
        # 空行
        story.append(Spacer(1, 10))
        
        # 每20题分页
        if q_num % 20 == 0:
            story.append(PageBreak())
        
        q_num += 1
    
    # 生成 PDF
    doc.build(story)
    
    return filepath


def get_export_formats():
    """获取支持的导出格式"""
    return [
        {'value': 'word', 'label': 'Word 文档 (.docx)', 'extension': '.docx'},
        {'value': 'pdf', 'label': 'PDF 文档 (.pdf)', 'extension': '.pdf'}
    ]


def get_export_options():
    """获取导出选项"""
    return {
        'include_answer': {
            'type': 'checkbox',
            'label': '包含答案',
            'default': True
        },
        'include_analysis': {
            'type': 'checkbox',
            'label': '包含解析（如有）',
            'default': True
        }
    }


def get_supported_types():
    """获取支持的题目类型"""
    return ['single', 'multiple', 'judge', 'fill', 'short_answer']
